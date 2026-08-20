#!/usr/bin/env python3
"""Actualiza la entrega Word de Aurora conservando su contenido historico."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
import re
import shutil
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TODAY_ISO = "2026-08-20"
TODAY_LONG = "20 de agosto de 2026"
VERSION = "1.3"
RESPONSIBLE = "Dirección Nacional de Defensoría Pública (DNDP) - Grupo de Transformación Digital"
APPROVAL = "Pendiente aprobación institucional"
MARKER = "Actualización técnica y funcional al 20 de agosto de 2026"

NAVY = "073763"
BLUE = "D9EAF7"
GREEN = "E7F3E8"
YELLOW = "FFF3D2"
GRAY = "F3F5F7"


UPDATES = {
    "ACCESO_REPOSITORIO_AURORA.docx": {
        "change": "Actualización de custodia, fuente vigente, estructura documental y trazabilidad del corte técnico.",
        "sections": [
            ("Estado vigente de la entrega", [
                "La versión documental consolidada corresponde al corte técnico e4ee2f7 de la rama master, sincronizado el 20 de agosto de 2026. SharePoint conserva la entrega inicial y Azure DevOps continúa definido como repositorio institucional de destino; cualquier espejo técnico adicional no reemplaza esa custodia.",
                "Los once documentos Word ubicados en la raíz y los cinco diagramas de diagramas_arquitectura/ forman la entrega documental 1.3. Las copias bajo documentacion/documentacion_word/ se conservan como antecedente histórico y no deben usarse como versión vigente.",
            ]),
            ("Procedimiento actualizado de entrega", [
                "Registrar commit, rama, fecha, responsable y resultado de validación antes de promover una versión.",
                "Entregar código, .env.example, Dockerfile, docker-compose.yml, documentación, diagramas y evidencias de pruebas sin incluir secretos, certificados, archivos de carga ni datos personales.",
                "Crear la etiqueta o mecanismo institucional de versión después de la aprobación y conservar el acta o evidencia de recepción.",
            ]),
        ],
        "diagrams": ["01_contexto_aurora.png"],
    },
    "ARQUITECTURA_SISTEMA_AURORA.docx": {
        "change": "Actualización de arquitectura física, autenticación, roles, módulos, persistencia auxiliar, reportes y PWA.",
        "sections": [
            ("Arquitectura lógica vigente", [
                "Aurora mantiene una arquitectura cliente-servidor modular. React 19 y Vite implementan la presentación; Node.js y Express 5 exponen la API /api, aplican autenticación y autorización, y sirven el frontend compilado; Oracle conserva los datos de negocio; Python procesa las cargas Excel hacia staging y ETL.",
                "La solución no se dividió en microservicios. El despliegue productivo sigue construyendo un único contenedor de aplicación, con módulos internos separados por rutas, servicios, dominios y repositorios.",
            ]),
            ("Módulos incorporados o ampliados después del corte de junio", [
                "Administración de usuarios autorizados, importación CSV y persistencia en el volumen aurora_auth_users.",
                "Autenticación alternativa LDAP/LDAPS y modo de acceso open o managed, además de Microsoft Entra ID y del administrador local temporal.",
                "Rol pag verificado en frontend y backend para asignación y reasignación de casos.",
                "Reporte de atenciones por defensor con filtros y generación PDF en el navegador.",
                "Catálogos versionados, homologación de centros, auditoría de normalización y tolerancia controlada a codificación heredada.",
                "Aviso de tratamiento de datos, manual interactivo con videos locales y PWA con cola offline ligada a la identidad autenticada.",
                "Depuración administrativa acotada de actuaciones ficticias y reparación automática del historial de cargas.",
            ]),
            ("Arquitectura física vigente", [
                "Para usuarios y equipos cliente, Aurora se publica exclusivamente por HTTPS en el puerto 443 y la URL institucional no incluye :7860. El puerto 7860 pertenece únicamente a la red interna de Docker, donde escucha Node.js dentro del contenedor; no es un puerto público ni debe habilitarse para acceso directo. TLS puede terminar en el proxy institucional o directamente en Aurora mediante HTTPS_KEY_PATH y HTTPS_CERT_PATH.",
                "Los volúmenes aurora_auth_users y aurora_cargas_bd separan el estado operativo del ciclo de vida del contenedor. Oracle continúa externo por el puerto de servicio configurado, normalmente 1521.",
            ]),
            ("Roles vigentes", [
                "user: consulta, formularios, historial, formatos y módulos ordinarios.",
                "pag: acceso adicional a asignación, reasignación y operaciones PAG protegidas.",
                "admin: administración del directorio interno de usuarios y soporte autorizado.",
                "admin, carguebd o cargas_bd: acceso al módulo de cargas cuando coinciden con CARGUEBD_ADMIN_ROLES.",
            ]),
        ],
        "diagrams": [
            "01_contexto_aurora.png", "02_componentes_aurora.png", "03_despliegue_aurora.png",
            "04_autenticacion_roles_aurora.png", "05_cargas_excel_staging_etl_aurora.png",
        ],
    },
    "DESCRIPCION_CODIGO_FUENTE_AURORA.docx": {
        "change": "Actualización del inventario de código, rutas, servicios, módulos frontend, variables y cobertura de pruebas.",
        "sections": [
            ("Estructura vigente del backend", [
                "backend/routes incluye auth, ppl, defensores, formatos, health, reportes, adminCargas y adminUsers.",
                "backend/services incluye autenticación, LDAP, directorio de usuarios, PPL, cargas, reportes, importación CSV, homologación y depuración de actuaciones.",
                "backend/repositories/oracle separa persona, situación, gestión jurídica, asignación, defensores, PAG, conducta, reportes y limpieza controlada.",
                "backend/catalogs y backend/domain incorporan catálogos versionados, homologación y derivación del estado del caso.",
            ]),
            ("Estructura vigente del frontend", [
                "Las vistas actuales son Inicio, Formulario de atención, Usuarios asignados, Reporte de atenciones, PAG - Asignación, Caja de herramientas, Manual interactivo, Cargas mensuales y Usuarios autorizados.",
                "DataTreatmentNotice exige aceptación antes de habilitar la aplicación. FEATURE_FLAGS controla módulos opcionales y el service worker conserva el shell y una cola offline acotada por identidad.",
                "ReporteAtencionesDefensores y reporteAtencionesPdf incorporan filtros, consulta y exportación PDF. AdminUsuarios administra el directorio interno sin almacenar contraseñas institucionales.",
            ]),
            ("Tecnologías y operación vigentes", [
                "Frontend: React 19.2, Vite 7, MSAL Browser, jsPDF, Vitest y ESLint.",
                "Backend: Node.js 20, Express 5, Helmet, CORS, JWT, JWKS, ldapts, Multer y node-oracledb.",
                "Cargas: Python 3, pandas, openpyxl y oracledb; despliegue principal mediante Docker Compose.",
            ]),
            ("Variables nuevas o ampliadas", [
                "AUTH_USER_ACCESS_MODE, AUTH_BOOTSTRAP_ADMIN_EMAILS, AUTH_USER_STORE_PATH, AUTH_USER_IMPORT_MAX_MB, AUTH_USER_IMPORT_MAX_ROWS y AUTH_USER_SYNC_REQUIRED.",
                "LDAP_ENABLED, LDAP_URL, LDAP_DOMAIN, LDAP_ALLOWED_EMAIL_DOMAINS y LDAP_TIMEOUT_MS.",
                "AURORA_VIDEOS_DIR, AURORA_CARGAS_TMP_DIR y variables de retención, reparación, limpieza y longitud de error del historial de cargas.",
            ]),
        ],
        "diagrams": ["02_componentes_aurora.png", "04_autenticacion_roles_aurora.png"],
    },
    "DESCRIPCION_MODELO_DATOS_AURORA.docx": {
        "change": "Restauración y actualización de referencias del modelo, persistencia auxiliar, transacciones y criterios de vigencia.",
        "sections": [
            ("Referencia y alcance", [
                "El MER y el diccionario aprobados por el DBA continúan siendo la definición contractual. Este documento describe el uso observado desde el código y no debe sustituir cardinalidades, nulabilidad, índices o restricciones certificadas.",
                "El catálogo local de formatos vigente es backend/data/formatos.js. No existe formatos.mock.js en la estructura actual.",
            ]),
            ("Objetos y relaciones vigentes", [
                "Las entidades de negocio observadas incluyen PERSONA, SITUACION_CARCELARIA, GESTION_JURIDICA, ASIGNACION, DEFENSORES, PAG, CALIFICACION_CONDUCTA y REGIONALES.",
                "PONAL, SISIPEC y AURORA_10 son staging para cargas externas; LOG_CARGA registra procedimientos ETL cuando el ambiente lo provee.",
                "ReporteAtencionesRepository consulta gestiones, situaciones, personas, asignaciones, defensores y regionales sin crear un modelo de persistencia independiente para el PDF.",
            ]),
            ("Criterios de vigencia y trazabilidad", [
                "La acción operativa se persiste en GESTION_JURIDICA.ACCION_REALIZAR. La fuente y fecha de corte provienen de SITUACION_CARCELARIA.FUENTE y FECHA_CORTE.",
                "Cuando existen varias situaciones, se prioriza la fecha efectiva más reciente y se usa ACTIVO como criterio de desempate. La interfaz informa fuente, corte y coexistencia de filas activas e inactivas.",
                "La auditoría nominal completa de cada edición sigue requiriendo columnas o una tabla transaccional aprobada por el DBA.",
            ]),
            ("Persistencia auxiliar e integridad", [
                "El directorio de usuarios autorizados se conserva en JSON dentro del volumen aurora_auth_users; guarda identidad administrativa, roles, estado y fechas, pero no contraseñas institucionales.",
                "El volumen aurora_cargas_bd conserva registro, archivos, logs y evidencias operativas. La depuración controlada escribe auditoría JSON Lines.",
                "Asignaciones, gestiones y limpiezas controladas usan transacciones. ORACLE_GESTION_ID_SEQUENCE es preferible frente al respaldo MAX(ID)+1 cuando el ambiente dispone de secuencia.",
            ]),
            ("Limitaciones y recomendaciones restauradas", [
                "Validar privilegios, índices, procedimientos y objetos reales con el DBA antes de producción.",
                "Mantener una matriz entre campos de interfaz, nombres canónicos de API y columnas Oracle.",
                "Tratar CSV y Excel con datos personales como información restringida; anonimizar pruebas y no incorporarlos a evidencias públicas.",
                "Ejecutar pruebas de escritura solamente contra un esquema temporal autorizado.",
            ]),
        ],
        "diagrams": ["05_cargas_excel_staging_etl_aurora.png"],
    },
    "GUIA_DESPLIEGUE_AURORA.docx": {
        "change": "Actualización de puertos, volúmenes, autenticación, variables, healthcheck y validación del despliegue vigente.",
        "sections": [
            ("Configuración productiva vigente", [
                "Aurora se publica exclusivamente por HTTPS en el puerto 443 para usuarios y equipos cliente. La URL institucional no debe incluir :7860. PORT=7860 identifica solamente el proceso Node.js dentro del contenedor y no debe exponerse en firewall, DNS, proxy ni documentación dirigida al usuario.",
                "Docker Compose crea los volúmenes aurora_auth_users y aurora_cargas_bd. No usar docker compose down -v durante una actualización ordinaria porque eliminaría el estado persistente.",
                "El healthcheck selecciona HTTP o HTTPS según HTTPS_KEY_PATH y HTTPS_CERT_PATH y consulta /api/health dentro del contenedor.",
            ]),
            ("Variables que deben revisarse en cada ambiente", [
                "Identidad: AZURE_AD_*, LDAP_*, AUTH_USER_ACCESS_MODE, AUTH_BOOTSTRAP_ADMIN_EMAILS y AUTH_USER_SYNC_REQUIRED.",
                "Persistencia: AUTH_USER_STORE_PATH, AURORA_CARGAS_DIR y AURORA_CARGAS_TMP_DIR.",
                "Cargas: límites de archivo, retención, cantidad de registros, reparación, limpieza controlada y procedimiento SISIPEC.",
                "Publicación: HOST_PORT, PORT, CORS_ORIGIN, HTTPS_KEY_PATH y HTTPS_CERT_PATH.",
            ]),
            ("Secuencia de actualización recomendada", [
                "Respaldar .env, directorio de usuarios, historial/logs de cargas y cualquier volumen administrado antes de actualizar.",
                "Ejecutar docker compose config --quiet y docker compose build aurora.",
                "Aplicar docker compose up -d --force-recreate aurora cuando cambie .env; usar --build cuando cambie código o dependencias.",
                "Validar docker compose ps, /api/health, /api/health/db y /api/auth/config; después ejecutar login, consulta, formulario, reportes, asignación, formatos y módulos administrativos según rol.",
            ]),
            ("Advertencia sobre referencias históricas", [
                "Las direcciones 172.31.64.7, localhost y los puertos 7860/8787 documentados previamente son evidencias o ejemplos de pruebas de mayo y junio. No constituyen la URL productiva vigente.",
                "PM2 se mantiene únicamente como alternativa Node.js directa aprobada por Infraestructura; no debe administrar en paralelo el proceso que ejecuta Docker Compose.",
            ]),
        ],
        "diagrams": ["03_despliegue_aurora.png"],
    },
    "INFRAESTRUCTURA_AURORA.docx": {
        "change": "Actualización del despliegue, puertos, volúmenes persistentes, TLS y dependencias institucionales.",
        "sections": [
            ("Topología vigente", [
                "El acceso institucional se realiza exclusivamente por HTTPS en el puerto público 443, mediante proxy/balanceador o TLS directo en Aurora. El puerto 7860 existe solo dentro de Docker para la comunicación con el proceso Node.js; no se publica ni se utiliza desde equipos cliente.",
                "El contenedor integra frontend, backend, Python y tutoriales. Se conecta a Oracle y a los servicios de identidad configurados; no expone Oracle al cliente.",
                "Los volúmenes aurora_auth_users y aurora_cargas_bd requieren respaldo, permisos y retención definidos. backend/certs se monta como solo lectura cuando se utiliza TLS directo.",
            ]),
            ("Responsabilidades actualizadas", [
                "Infraestructura/TICS: DNS, proxy, certificado, firewall, Docker, respaldos de volúmenes, monitoreo y publicación.",
                "DBA/TICS: servicio Oracle, esquema, privilegios, secuencias, procedimientos ETL, índices y respaldo de base de datos.",
                "Administrador de identidad: App Registration, redirect URI, grupos, roles, LDAP/LDAPS y cuentas autorizadas.",
                "Equipo DNDP: código, configuración documentada, pruebas, soporte funcional y actualización de esta entrega.",
            ]),
            ("Disponibilidad y monitoreo", [
                "/api/health mide disponibilidad del proceso; /api/health/db diferencia disponibilidad de Oracle.",
                "La arquitectura actual no incluye balanceo interno, réplica de aplicación, almacenamiento distribuido ni alta disponibilidad de Oracle. Cualquier ampliación debe diseñarse y aprobarse explícitamente.",
            ]),
        ],
        "diagrams": ["03_despliegue_aurora.png"],
    },
    "LINEAMIENTOS_SEGURIDAD_AURORA.docx": {
        "change": "Actualización de proveedores de identidad, autorización por rol, usuarios administrados, PWA y protección de datos.",
        "sections": [
            ("Identidad y autorización vigentes", [
                "Microsoft Entra ID es el proveedor institucional principal. LDAP/LDAPS puede habilitarse como alternativa de bind directo y el administrador local solo se admite como mecanismo temporal controlado.",
                "El directorio interno permite modo open o managed. En managed, una identidad válida también debe estar habilitada internamente. La aplicación no almacena la contraseña institucional.",
                "Los roles user, pag, admin y los roles de cargas se validan en backend. Ocultar una opción en React no constituye autorización.",
            ]),
            ("Protección de datos y sesión", [
                "El usuario debe aceptar el aviso de tratamiento de datos antes de operar los módulos. Las sesiones usan JWT interno con secreto, emisor, audiencia y expiración configurables.",
                "La cola offline de la PWA está acotada y ligada a la identidad; cerrar sesión o cambiar de cuenta descarta operaciones de la identidad anterior.",
                "Los errores públicos se sanitizan. Logs, reportes y evidencias no deben exponer contraseñas, tokens, cadenas de conexión ni datos personales innecesarios.",
            ]),
            ("Custodia técnica", [
                ".env, certificados, llaves privadas, CSV/Excel reales, almacenamiento de usuarios y volúmenes de cargas deben quedar fuera de Git y con permisos mínimos.",
                "Los tutoriales incluidos en la imagen se validan con SHA256SUMS durante el build. Las cargas aceptan .xlsx, aplican límites configurables y registran usuario, estado y log.",
                "La depuración de actuaciones exige rol autorizado, confirmación explícita, defensor exacto configurado, conteo previo, transacción y bitácora.",
            ]),
        ],
        "diagrams": ["04_autenticacion_roles_aurora.png"],
    },
    "LINEA_BASE_CASO_NEGOCIO_AURORA.docx": {
        "change": "Actualización del alcance funcional, gobierno de usuarios, reportes, calidad de datos y operación institucional.",
        "sections": [
            ("Evolución funcional consolidada", [
                "La arquitectura web mantiene la centralización progresiva de los aprendizajes de Aurora, Celeste y MARIA. El alcance actual fortalece condenados y sindicados, asignación PAG, formularios, historial, formatos y cargas.",
                "Se incorporaron reporte de atenciones por defensor, administración de usuarios autorizados, tratamiento de datos, catálogos homologados y tutoriales locales.",
            ]),
            ("Beneficios adicionales", [
                "Gobierno de acceso mediante roles y directorio interno administrado, sin administrar contraseñas institucionales.",
                "Mejor seguimiento mediante reportes PDF, fuente y fecha de corte visibles, estados derivados y filtros homologados.",
                "Mayor continuidad operativa mediante despliegue reproducible, volúmenes persistentes, pruebas automatizadas y reparación controlada de registros de cargas.",
                "Mejor calidad de datos mediante catálogos versionados, auditoría de homologación y normalización tolerante de información heredada.",
            ]),
            ("Restricciones vigentes", [
                "La publicación, alta disponibilidad, identidad, Oracle y ETL dependen de las áreas institucionales responsables.",
                "La auditoría nominal completa por campo requiere evolución aprobada del modelo de datos.",
                "La centralización funcional continúa siendo gradual; las reglas jurídicas y diferencias poblacionales deben validarse con los responsables funcionales.",
            ]),
        ],
        "diagrams": ["01_contexto_aurora.png"],
    },
    "MANUAL_TECNICO_AURORA.docx": {
        "change": "Actualización integral de módulos, API, autenticación, despliegue, datos, reportes y pruebas vigentes.",
        "sections": [
            ("Delta técnico posterior a junio", [
                "La arquitectura continúa en un solo contenedor, pero el backend amplió sus capas de dominio, servicios y repositorios. Se añadieron usuarios autorizados, LDAP, rol PAG, reportes, homologación, limpieza controlada y nuevos contratos de pruebas.",
                "El frontend añadió ReporteAtencionesDefensores, AdminUsuarios, PagUsuariosAdmin, DataTreatmentNotice, catálogos auxiliares y utilidades de PDF, asignación, estado y validación.",
            ]),
            ("API incorporada o ampliada", [
                "GET /api/reportes/atenciones-defensores/opciones y GET /api/reportes/atenciones-defensores.",
                "GET, POST, PATCH y DELETE bajo /api/admin/users, además de previsualización e importación CSV.",
                "POST /api/ppl/desasignar-defensor y auditoría de homologación para operaciones PAG autorizadas.",
                "Operaciones de previsualización y limpieza acotada de actuaciones dentro de /api/admin/cargas.",
            ]),
            ("Configuración y despliegue vigente", [
                "HOST_PORT=443 es el único puerto público para usuarios. PORT=7860 es exclusivamente interno al contenedor, no aparece en la URL institucional y no debe exponerse directamente. ENABLE_STARTUP_WARMUP está activo por defecto.",
                "El despliegue incorpora volúmenes para usuarios y cargas, entrypoint con reducción de privilegios, tutoriales verificados y healthcheck sensible a HTTP/HTTPS.",
                "Las nuevas familias de variables son AUTH_USER_*, LDAP_*, AURORA_VIDEOS_DIR, AURORA_CARGAS_TMP_DIR y controles de retención/reparación de cargas.",
            ]),
            ("Datos y reglas vigentes", [
                "Los listados y filtros usan catálogos homologados, paginación acotada, caché temporal y normalización de codificación heredada.",
                "Asignar, reasignar y desasignar conserva reglas de vigencia e historial. Fuente y fecha de corte de SITUACION_CARCELARIA se muestran para sindicados y condenados cuando están disponibles.",
                "Los reportes consultan Oracle mediante repositorio y servicio dedicados y se convierten a PDF en frontend con jsPDF.",
            ]),
            ("Validación automatizada del corte", [
                "El 20 de agosto de 2026 se ejecutó npm run qa:smoke con resultado satisfactorio: lint, 145 pruebas frontend en 12 archivos, build de producción y 14 suites backend.",
                "El build emitió una advertencia no bloqueante por un chunk JavaScript superior a 500 kB. Las pruebas de integración con Oracle, escritura controlada y validación funcional manual requieren ambiente autorizado y no se consideran ejecutadas por este resultado.",
            ]),
        ],
        "diagrams": ["02_componentes_aurora.png", "03_despliegue_aurora.png", "04_autenticacion_roles_aurora.png"],
    },
    "MANUAL_USUARIO_AURORA.docx": {
        "change": "Actualización de acceso, tratamiento de datos, roles, reportes, usuarios autorizados, cargas y operación segura.",
        "sections": [
            ("Ingreso y aceptación de condiciones", [
                "Después de autenticar la cuenta institucional, Aurora presenta el aviso de tratamiento de datos. Es necesario aceptarlo para habilitar la navegación; rechazarlo cierra la sesión.",
                "Las opciones visibles dependen de los roles. Un acceso denegado debe reportarse para validar cuenta, estado y roles; no se debe usar la cuenta de otra persona.",
            ]),
            ("Módulos vigentes", [
                "Formulario de atención y Usuarios asignados: consulta, filtros, historial, fuente, fecha de corte, diligenciamiento y guardado.",
                "Descargar reporte de atención: filtra atenciones por criterios disponibles y genera un PDF para el uso institucional autorizado.",
                "PAG - Asignación de casos: solo para rol pag; permite asignar, reasignar y desasignar conforme a las reglas visibles.",
                "Caja de herramientas y Manual interactivo: formatos y tres tutoriales locales para condenados ERON, Ley 906/CDT y PAG.",
                "Cargas mensuales: para roles autorizados; recibe PONAL, SISIPEC o Aurora 1.0, muestra estado/log y permite reintentos controlados.",
                "Usuarios autorizados: para admin; habilita, deshabilita e importa usuarios y roles, pero no administra contraseñas institucionales.",
            ]),
            ("Trabajo sin conexión y buenas prácticas", [
                "La aplicación instalada puede conservar el shell y encolar determinadas escrituras cuando falla la red. Verifique siempre el mensaje final de guardado al recuperar conectividad.",
                "Cerrar sesión o cambiar de cuenta descarta operaciones pendientes de la identidad anterior. No use el modo sin conexión en equipos compartidos para información sensible.",
                "No comparta capturas o reportes con datos personales por canales no autorizados. Registre URL, hora, módulo y mensaje al reportar un incidente, sin enviar contraseñas ni tokens.",
            ]),
        ],
        "diagrams": [],
    },
    "PRUEBAS_FUNCIONALES_AURORA.docx": {
        "change": "Actualización de matriz, cobertura automatizada y separación entre resultados comprobados y pruebas pendientes de ambiente.",
        "sections": [
            ("Resultado automatizado del 20 de agosto de 2026", [
                "Comando ejecutado: npm run qa:smoke.",
                "Frontend lint: aprobado.",
                "Frontend Vitest: 12 archivos y 145 pruebas aprobadas.",
                "Frontend build: aprobado; PWA inyectó seis assets. Se conserva advertencia no bloqueante por un chunk superior a 500 kB.",
                "Backend: 14 suites aprobadas, incluyendo autenticación, PAG, usuarios CSV, cargas, limpieza, asignación, fechas, codificación, catálogos, auditoría, filtros, contratos y reportes.",
            ]),
            ("Cobertura añadida desde junio", [
                "Acceso PAG y seguridad de asignaciones; importación CSV y directorio de usuarios.",
                "Reparación y limpieza del historial de cargas; limpieza acotada de actuaciones.",
                "Homologación de catálogos, auditoría, normalización de texto y estados de condenados.",
                "Contrato de Usuarios asignados, filtros, fechas de insistencia y reporte de atenciones.",
                "Reglas de formularios, PWA, PDF, manual interactivo y asignación en frontend.",
            ]),
            ("Pruebas que continúan pendientes", [
                "La matriz manual AUR-001 a AUR-012 conserva su estado histórico hasta contar con acta y evidencia del ambiente institucional.",
                "No se ejecutaron en esta actualización smoke Oracle, integraciones Oracle, regresión API contra servidor, escritura controlada ni los casos funcionales manuales con usuarios reales.",
                "Antes de aprobar producción deben validarse HTTPS, Entra ID/LDAP, roles user/pag/admin/cargas, Oracle, ETL, reportes PDF, formularios, asignación, formatos, usuarios autorizados y respaldo de volúmenes.",
            ]),
        ],
        "diagrams": ["05_cargas_excel_staging_etl_aurora.png"],
    },
}


REPLACEMENTS = {
    "formatos.mock.js": "formatos.js",
    "docs/16_cargas_staging_etl_bd.md": "scripts/cargas_bd/README.md",
    "Docker Compose publica ${HOST_PORT:-443}:${PORT:-7860}. El puerto 7860 es interno al contenedor y el puerto 443 es el valor externo predeterminado. TLS puede terminar en el proxy institucional o directamente en Aurora mediante HTTPS_KEY_PATH y HTTPS_CERT_PATH.":
        "Para usuarios y equipos cliente, Aurora se publica exclusivamente por HTTPS en el puerto 443 y la URL institucional no incluye :7860. El puerto 7860 pertenece únicamente a la red interna de Docker, donde escucha Node.js dentro del contenedor; no es un puerto público ni debe habilitarse para acceso directo. TLS puede terminar en el proxy institucional o directamente en Aurora mediante HTTPS_KEY_PATH y HTTPS_CERT_PATH.",
    "El mapeo predeterminado es HOST_PORT=443 hacia PORT=7860. Un puerto alterno puede usarse para diagnóstico, pero la URL institucional definitiva debe operar sobre HTTPS.":
        "Aurora se publica exclusivamente por HTTPS en el puerto 443 para usuarios y equipos cliente. La URL institucional no debe incluir :7860. PORT=7860 identifica solamente el proceso Node.js dentro del contenedor y no debe exponerse en firewall, DNS, proxy ni documentación dirigida al usuario.",
    "El acceso institucional usa HTTPS hacia un proxy/balanceador o directamente hacia Aurora. Docker publica por defecto 443 y entrega el tráfico al puerto interno 7860 del contenedor.":
        "El acceso institucional se realiza exclusivamente por HTTPS en el puerto público 443, mediante proxy/balanceador o TLS directo en Aurora. El puerto 7860 existe solo dentro de Docker para la comunicación con el proceso Node.js; no se publica ni se utiliza desde equipos cliente.",
    "HOST_PORT usa 443 por defecto y PORT 7860 dentro del contenedor. ENABLE_STARTUP_WARMUP está activo por defecto.":
        "HOST_PORT=443 es el único puerto público para usuarios. PORT=7860 es exclusivamente interno al contenedor, no aparece en la URL institucional y no debe exponerse directamente. ENABLE_STARTUP_WARMUP está activo por defecto.",
}


def font(size: int, bold: bool = False):
    candidates = [
        "/usr/share/fonts/dejavu-sans-fonts/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/dejavu-sans-fonts/DejaVuSans.ttf",
        "/usr/share/fonts/liberation-sans/LiberationSans-Bold.ttf" if bold else "/usr/share/fonts/liberation-sans/LiberationSans-Regular.ttf",
        "/usr/share/fonts/cantarell/Cantarell-Bold.otf" if bold else "/usr/share/fonts/cantarell/Cantarell-Regular.otf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def wrapped_lines(text: str, width: int):
    return textwrap.wrap(text, width=width, break_long_words=False, break_on_hyphens=False)


def draw_box(draw, xy, title, lines, fill=BLUE, title_size=27, body_size=20):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=10, fill=f"#{fill}", outline=f"#{NAVY}", width=2)
    draw.text((x1 + 18, y1 + 14), title, fill=f"#{NAVY}", font=font(title_size, True))
    draw.line((x1 + 18, y1 + 50, x2 - 18, y1 + 50), fill="#9FB5C9", width=1)
    y = y1 + 65
    for line in lines:
        for part in wrapped_lines(line, max(16, int((x2 - x1) / (body_size * 0.58)))):
            draw.text((x1 + 18, y), part, fill="#17283A", font=font(body_size))
            y += body_size + 7
        y += 2


def arrow(draw, start, end, label=None):
    draw.line((*start, *end), fill=f"#{NAVY}", width=5)
    x2, y2 = end
    x1, y1 = start
    angle_horizontal = abs(x2 - x1) >= abs(y2 - y1)
    if angle_horizontal:
        direction = 1 if x2 > x1 else -1
        head = [(x2, y2), (x2 - 16 * direction, y2 - 10), (x2 - 16 * direction, y2 + 10)]
    else:
        direction = 1 if y2 > y1 else -1
        head = [(x2, y2), (x2 - 10, y2 - 16 * direction), (x2 + 10, y2 - 16 * direction)]
    draw.polygon(head, fill=f"#{NAVY}")
    if label:
        mx, my = (x1 + x2) // 2, (y1 + y2) // 2
        bbox = draw.textbbox((0, 0), label, font=font(16))
        w, h = bbox[2] - bbox[0] + 14, bbox[3] - bbox[1] + 8
        draw.rounded_rectangle((mx - w // 2, my - h // 2, mx + w // 2, my + h // 2), 5, fill="white", outline="#B8C8D8")
        draw.text((mx - w // 2 + 7, my - h // 2 + 3), label, fill="#17283A", font=font(16))


def canvas(title, subtitle):
    image = Image.new("RGB", (1800, 1000), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, 1800, 112), fill=f"#{NAVY}")
    draw.text((42, 18), title, fill="white", font=font(42, True))
    draw.text((44, 72), subtitle, fill="white", font=font(22))
    return image, draw


def diagram_context(path):
    image, draw = canvas("Diagrama de contexto - Aurora", "Corte técnico vigente al 20 de agosto de 2026")
    draw_box(draw, (55, 180, 350, 350), "Usuarios", ["Funcional (user)", "PAG", "Administrador"], BLUE)
    draw_box(draw, (485, 180, 820, 350), "Navegador / PWA", ["React 19 + Vite", "Aviso de datos", "Módulos por rol"], GRAY)
    draw_box(draw, (955, 150, 1325, 390), "Sistema Aurora", ["Express 5 / API", "JWT interno", "Servicios y repositorios", "Reportes y cargas"], BLUE)
    draw_box(draw, (1460, 150, 1740, 340), "Identidad", ["Microsoft Entra ID", "LDAP/LDAPS", "Local temporal"], GRAY)
    draw_box(draw, (1460, 470, 1740, 655), "Oracle", ["Personas", "Gestiones", "Asignaciones", "Staging / ETL"], GREEN)
    draw_box(draw, (955, 510, 1325, 710), "Estado persistente", ["Usuarios autorizados", "Archivos y logs", "Volúmenes Docker"], YELLOW)
    draw_box(draw, (290, 690, 650, 875), "SharePoint", ["Entrega inicial", "Versión cero"], YELLOW)
    draw_box(draw, (700, 690, 1060, 875), "Azure DevOps", ["Repositorio institucional", "Control posterior"], YELLOW)
    arrow(draw, (350, 265), (485, 265), "HTTPS")
    arrow(draw, (820, 265), (955, 265), "/api")
    arrow(draw, (1325, 230), (1460, 230), "login")
    arrow(draw, (1460, 300), (1325, 300), "token")
    arrow(draw, (1325, 350), (1460, 520), "SQL")
    arrow(draw, (1140, 390), (1140, 510), "estado")
    arrow(draw, (650, 780), (700, 780), "migración")
    draw.text((55, 930), "La autenticación externa no sustituye la autorización interna por roles y estado de usuario.", fill="#17283A", font=font(20, True))
    image.save(path)


def diagram_components(path):
    image, draw = canvas("Diagrama de componentes - Aurora", "Arquitectura modular dentro de un único servicio desplegable")
    draw_box(draw, (50, 150, 430, 390), "Presentación", ["React / Vite / PWA", "Formulario y listados", "PAG y reportes PDF", "Administración y manual"], BLUE)
    draw_box(draw, (520, 150, 900, 390), "API y seguridad", ["Express /api", "Helmet y CORS", "JWT y roles", "Entra ID / LDAP"], GRAY)
    draw_box(draw, (990, 150, 1370, 390), "Servicios y dominio", ["PPL y estado", "Usuarios y CSV", "Cargas y limpieza", "Reportes / catálogos"], BLUE)
    draw_box(draw, (1460, 150, 1750, 390), "Repositorios", ["SQL parametrizado", "Pool Oracle", "Normalización", "Transacciones"], GREEN)
    draw_box(draw, (260, 560, 650, 810), "Procesamiento Python", ["PONAL / SISIPEC", "Aurora 1.0", "Validación Excel", "Staging y ETL"], YELLOW)
    draw_box(draw, (760, 560, 1150, 810), "Persistencia auxiliar", ["aurora_auth_users", "aurora_cargas_bd", "Logs y auditoría", "Tutoriales verificados"], YELLOW)
    draw_box(draw, (1260, 560, 1650, 810), "Oracle", ["Modelo de negocio", "Catálogos", "Staging", "Procedimientos ETL"], GREEN)
    arrow(draw, (430, 270), (520, 270), "HTTP")
    arrow(draw, (900, 270), (990, 270), "rutas")
    arrow(draw, (1370, 270), (1460, 270), "datos")
    arrow(draw, (1180, 390), (500, 560), "carga")
    arrow(draw, (650, 685), (760, 685), "logs")
    arrow(draw, (1150, 685), (1260, 685), "SQL / ETL")
    draw.text((50, 920), "Las capas separan responsabilidades; el artefacto de despliegue continúa siendo un contenedor Aurora.", fill="#17283A", font=font(20, True))
    image.save(path)


def diagram_deployment(path):
    image, draw = canvas("Diagrama de despliegue - Aurora", "Publicación HTTPS, contenedor único y dependencias externas")
    draw_box(draw, (45, 260, 315, 500), "Cliente", ["Navegador institucional", "Red autorizada / VPN", "HTTPS"], GRAY)
    draw_box(draw, (400, 230, 710, 530), "Publicación segura", ["HTTPS público 443", "Proxy / balanceador", "o TLS directo", "Único acceso de usuarios"], BLUE)
    draw_box(draw, (820, 150, 1280, 650), "Servidor de aplicaciones", ["Linux + Docker Compose", "Contenedor Aurora", "Frontend + backend + Python", "Puerto 7860 solo interno", "No accesible a usuarios", "Healthcheck /api/health"], BLUE)
    draw_box(draw, (1390, 190, 1750, 430), "Identidad", ["Microsoft Entra ID", "LDAP/LDAPS", "Conectividad controlada"], GRAY)
    draw_box(draw, (1390, 540, 1750, 790), "Base de datos", ["Oracle", "Servicio habitual 1521", "Esquema y permisos DBA", "Staging / ETL"], GREEN)
    draw_box(draw, (700, 735, 1080, 925), "Volúmenes", ["aurora_auth_users", "aurora_cargas_bd", "Respaldo y retención"], YELLOW)
    draw_box(draw, (1120, 735, 1360, 925), "Certificados", ["Montaje solo lectura", "Fuera de Git"], YELLOW)
    arrow(draw, (315, 380), (400, 380), "HTTPS")
    arrow(draw, (710, 380), (820, 380), "red Docker")
    arrow(draw, (1280, 270), (1390, 270), "identidad")
    arrow(draw, (1280, 560), (1390, 650), "Oracle")
    arrow(draw, (950, 650), (900, 735), "estado")
    arrow(draw, (1170, 650), (1240, 735), "TLS")
    image.save(path)


def diagram_auth(path):
    image, draw = canvas("Diagrama de autenticación y roles - Aurora", "Identidad externa, usuario administrado y autorización backend")
    draw_box(draw, (45, 170, 330, 390), "Usuario", ["Acceso por navegador", "Cuenta nominal"], BLUE)
    draw_box(draw, (420, 135, 770, 425), "Proveedor de identidad", ["Microsoft Entra ID", "LDAP/LDAPS", "Administrador local temporal"], GRAY)
    draw_box(draw, (860, 135, 1210, 425), "Servicio de autenticación", ["Valida token o bind", "Tenant / dominio", "Límite de intentos", "Normaliza roles"], BLUE)
    draw_box(draw, (1300, 135, 1750, 425), "Directorio interno", ["Modo open o managed", "Usuario habilitado", "Roles administrativos", "Sin contraseñas"], YELLOW)
    draw_box(draw, (690, 550, 1110, 760), "JWT Aurora", ["Emisor y audiencia", "Expiración", "Identidad y roles"], BLUE)
    draw_box(draw, (80, 810, 410, 955), "user", ["Consulta y formularios"], GREEN)
    draw_box(draw, (500, 810, 830, 955), "pag", ["Asignación protegida"], GREEN)
    draw_box(draw, (920, 810, 1250, 955), "admin", ["Usuarios autorizados"], YELLOW)
    draw_box(draw, (1340, 810, 1720, 955), "cargas", ["admin / carguebd / cargas_bd"], YELLOW)
    arrow(draw, (330, 280), (420, 280), "login")
    arrow(draw, (770, 280), (860, 280), "token / bind")
    arrow(draw, (1210, 280), (1300, 280), "acceso")
    arrow(draw, (1520, 425), (1020, 550), "roles")
    arrow(draw, (900, 760), (245, 810))
    arrow(draw, (900, 760), (665, 810))
    arrow(draw, (900, 760), (1085, 810))
    arrow(draw, (900, 760), (1530, 810))
    image.save(path)


def diagram_loads(path):
    image, draw = canvas("Diagrama de cargas Excel / staging / ETL - Aurora", "Flujo administrativo con trazabilidad, validación y persistencia")
    draw_box(draw, (45, 170, 330, 430), "Fuentes .xlsx", ["PONAL", "SISIPEC", "Aurora 1.0"], BLUE)
    draw_box(draw, (420, 150, 760, 450), "Módulo de cargas", ["Rol autorizado", "Valida tipo y tamaño", "Directorio temporal", "Registro de usuario"], YELLOW)
    draw_box(draw, (850, 150, 1190, 450), "Servicio Python", ["Lectura y normalización", "Carga por lotes", "Errores sanitizados", "Log de ejecución"], YELLOW)
    draw_box(draw, (1280, 150, 1740, 450), "Oracle staging", ["PONAL / SISIPEC / AURORA_10", "Procedimientos ETL", "LOG_CARGA cuando aplica"], GREEN)
    draw_box(draw, (1280, 610, 1740, 865), "Tablas de negocio", ["PERSONA", "SITUACION_CARCELARIA", "GESTION_JURIDICA", "ASIGNACION"], GREEN)
    draw_box(draw, (720, 610, 1110, 865), "Historial operativo", ["Estado y responsable", "Archivos y logs", "Retención configurable", "Reparación y respaldo"], GRAY)
    draw_box(draw, (130, 610, 560, 865), "Controles", ["Reintento autorizado", "Previsualización", "Limpieza acotada", "Auditoría JSON Lines"], GRAY)
    arrow(draw, (330, 290), (420, 290), "archivo")
    arrow(draw, (760, 290), (850, 290), "trabajo")
    arrow(draw, (1190, 290), (1280, 290), "insert")
    arrow(draw, (1510, 450), (1510, 610), "ETL")
    arrow(draw, (1020, 450), (920, 610), "estado / log")
    arrow(draw, (620, 450), (345, 610), "seguimiento")
    draw.text((45, 930), "Las pruebas y cargas reales requieren ambiente autorizado, respaldo y validación del DBA.", fill="#17283A", font=font(20, True))
    image.save(path)


def generate_diagrams():
    target = ROOT / "diagramas_arquitectura"
    target.mkdir(exist_ok=True)
    builders = {
        "01_contexto_aurora.png": diagram_context,
        "02_componentes_aurora.png": diagram_components,
        "03_despliegue_aurora.png": diagram_deployment,
        "04_autenticacion_roles_aurora.png": diagram_auth,
        "05_cargas_excel_staging_etl_aurora.png": diagram_loads,
    }
    for name, builder in builders.items():
        builder(target / name)
    historical_target = ROOT / "documentacion" / "diagramas_arquitectura"
    historical_target.mkdir(parents=True, exist_ok=True)
    for name in builders:
        shutil.copy2(target / name, historical_target / name)


def paragraph_texts(doc):
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(p.text for p in cell.paragraphs)
    return [text for text in texts if text.strip()]


def replace_in_paragraph(paragraph, replacements):
    for run in paragraph.runs:
        for old, new in replacements.items():
            if old in run.text:
                run.text = run.text.replace(old, new)


def replace_known_text(doc):
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, REPLACEMENTS)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, REPLACEMENTS)


def add_status(doc):
    status = f"Estado documental: vigente al {TODAY_LONG}. Versión {VERSION}."
    if any(status in p.text for p in doc.paragraphs):
        return
    for paragraph in doc.paragraphs:
        if "Control de cambios" in paragraph.text:
            inserted = paragraph.insert_paragraph_before(status)
            if inserted.runs:
                inserted.runs[0].bold = True
                inserted.runs[0].font.color.rgb = RGBColor(7, 55, 99)
            return
    paragraph = doc.paragraphs[0].insert_paragraph_before(status)
    paragraph.runs[0].bold = True


def update_control_table(doc, description):
    for table in doc.tables:
        if not table.rows:
            continue
        headers = [re.sub(r"\s+", " ", cell.text).strip().lower() for cell in table.rows[0].cells]
        if not ({"versión", "fecha", "responsable"} <= set(headers)):
            continue
        if any(VERSION in row.cells[0].text and TODAY_ISO in row.cells[1].text for row in table.rows[1:]):
            return
        row = table.add_row().cells
        values = [VERSION, TODAY_ISO, RESPONSIBLE, description, APPROVAL]
        for index, value in enumerate(values[: len(row)]):
            row[index].text = value
        return


def heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(7, 55, 99)
    return paragraph


def add_section_content(doc, title, items):
    heading(doc, title, 2)
    style_names = {style.name for style in doc.styles}
    for index, item in enumerate(items):
        is_bullet = len(items) >= 3 or item.startswith(("user:", "pag:", "admin:"))
        if is_bullet and "List Bullet" in style_names:
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.add_run(item)
        elif is_bullet:
            paragraph = doc.add_paragraph(style="List Paragraph" if "List Paragraph" in style_names else None)
            paragraph.add_run("• " + item)
        else:
            paragraph = doc.add_paragraph()
            paragraph.add_run(item)
        paragraph.paragraph_format.space_after = Pt(5)


def set_update_fields(doc):
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def append_update(doc, config):
    if any(MARKER in p.text for p in doc.paragraphs):
        return
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    heading(doc, MARKER, 1)
    intro = doc.add_paragraph()
    intro.add_run(
        "Este capítulo conserva como antecedente todo el contenido anterior y establece la interpretación vigente cuando exista una diferencia con secciones históricas. "
        "El corte corresponde a la rama master en e4ee2f7."
    )
    for title, items in config["sections"]:
        add_section_content(doc, title, items)
    if config["diagrams"]:
        heading(doc, "Diagramas vigentes", 2)
        for name in config["diagrams"]:
            image_path = ROOT / "diagramas_arquitectura" / name
            paragraph = doc.add_paragraph()
            paragraph.alignment = 1
            paragraph.add_run().add_picture(str(image_path), width=Inches(6.8))
            caption = doc.add_paragraph(f"Figura actualizada al {TODAY_LONG}: {name.replace('_', ' ').replace('.png', '')}.")
            caption.alignment = 1
            if caption.runs:
                caption.runs[0].italic = True


def refresh_appended_diagrams(doc, config):
    """Reemplaza los diagramas del apéndice sin tocar imágenes históricas."""
    diagram_names = config["diagrams"]
    if not diagram_names:
        return
    appended_shapes = list(doc.inline_shapes)[-len(diagram_names):]
    if len(appended_shapes) != len(diagram_names):
        raise RuntimeError("No fue posible localizar todos los diagramas añadidos")
    for shape, name in zip(appended_shapes, diagram_names):
        relation_id = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
        image_part = doc.part.related_parts[relation_id]
        image_part._blob = (ROOT / "diagramas_arquitectura" / name).read_bytes()


def update_docx(path, config):
    before = Document(path)
    before_texts = paragraph_texts(before)
    before_tables = len(before.tables)
    before_shapes = len(before.inline_shapes)

    replace_known_text(before)
    add_status(before)
    update_control_table(before, config["change"])
    append_update(before, config)
    refresh_appended_diagrams(before, config)
    before.core_properties.modified = datetime(2026, 8, 20, 12, 0, 0)
    before.core_properties.comments = f"Actualizado al {TODAY_LONG}; contenido histórico conservado."
    set_update_fields(before)
    before.save(path)

    after = Document(path)
    after_text = "\n".join(paragraph_texts(after))
    missing = []
    for original in before_texts:
        expected = original
        for old, new in REPLACEMENTS.items():
            expected = expected.replace(old, new)
        if expected not in after_text:
            missing.append(expected[:120])
    if missing:
        raise RuntimeError(f"{path.name}: se perdieron {len(missing)} bloques de texto; ejemplo: {missing[0]}")
    if len(after.tables) < before_tables:
        raise RuntimeError(f"{path.name}: disminuyó el número de tablas")
    if len(after.inline_shapes) < before_shapes:
        raise RuntimeError(f"{path.name}: disminuyó el número de imágenes")
    if not any(MARKER in p.text for p in after.paragraphs):
        raise RuntimeError(f"{path.name}: no se encontró el marcador de actualización")
    return {
        "paragraphs_before": len(before_texts),
        "paragraphs_after": len(paragraph_texts(after)),
        "tables_before": before_tables,
        "tables_after": len(after.tables),
        "images_before": before_shapes,
        "images_after": len(after.inline_shapes),
    }


def main():
    generate_diagrams()
    print(f"Actualizando {len(UPDATES)} documentos al {TODAY_ISO}...")
    for name, config in UPDATES.items():
        path = ROOT / name
        if not path.exists():
            raise FileNotFoundError(path)
        stats = update_docx(path, config)
        print(
            f"OK {name}: párrafos {stats['paragraphs_before']}->{stats['paragraphs_after']}, "
            f"tablas {stats['tables_before']}->{stats['tables_after']}, "
            f"imágenes {stats['images_before']}->{stats['images_after']}"
        )


if __name__ == "__main__":
    main()
